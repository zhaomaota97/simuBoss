import base64
import io

from docx import Document


def parse_word_base64(base64_payload: str) -> str:
    raw = base64.b64decode(base64_payload)
    document = Document(io.BytesIO(raw))

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))

    return "\n".join([*paragraphs, *table_rows]).strip()
